import re

import pandas as pd
from docx import *

####### VV CODE FOR MAIN SECTION OF MOLLAN VV #########

document = Document('Mollan_Irish_National_Inventory_of_Historic_Scientific_Instruments_1995.docx')
bolds = []
for para in document.paragraphs:
    # print('++++++++++++',para.text,'++++++++++++')
    # if found:
    #     temp_bolds_info = bolds_info.get(latest, [])
    #     temp_bolds_info.append(para.text)
    #     bolds_info[latest] = temp_bolds_info
    for run in para.runs:
        # if run.italic and run.text != 'IRISH NATIONAL INVENTORY OF HISTORIC SCIENTIFIC INSTRUMENTS':
        #     italics.append(run.text)
        if run.bold and run.text != 'IRISH NATIONAL INVENTORY OF HISTORIC SCIENTIFIC INSTRUMENTS' and run.text != 'AURORA' and run.text != 'COULOMB':
            # temp_bolds_info = bolds_info.get(latest, [])
            # temp_bolds_info = temp_bolds_info[:-1]
            # bolds_info[latest] = temp_bolds_info

            bolds.append(run.text)
            # found = True
            # counter = 0
            # latest = run.text
#

# boltalic_Dict={'bold_phrases':bolds}
bolds = bolds[30:5199]
# print(bolds)
# # print(bolds_info['3985 ARM059 DIVIDED CIRCLE WITH REVOLVING DISC'])
# # print(bolds_info['3950 ARM024 ALIDADE'])
# # print(bolds_info['3707 ARM020 CHRONOGRAPH'])
remove = ['CLONGOWES WOOD COLLEGE - CWC', 'Note: An overview of the collection, with some photographs, is given in ',
          'The Clongownian', ' - see Mollan 1987.', "ST PATRICK'S COLLEGE, MAYNOOTH - MAY",
          'Note: A comprehensive catalogue of the Maynooth instruments was',
          'published in 1994, and this gives fuller descriptions of the instruments, with photographs; it also includes more 20th Century items - see bibliography under Mollan 1994.',
          "NORTH MONASTERY CHRISTIAN BROTHERS' SCHOOL - NMC",
          'Note: This collection was sold in 1993. Some of the key instruments were acquired by the National Museum in Dublin, and the remainder sold to a London dealer. Those transferred to the National Museum are included also in its listing in this Inventory, and Irish-signed instruments sold to London are included also in the Out-of-Ireland listing.',
          'An overview of the collection is given Mollan 1993.', 'NATIONAL MUSEUM, DUBLIN - NMD',
          'Note: During the period of the compiling of the National Inventory, '
          'the collection at the National Museum has been dramatically increased and '
          'improved, most notably by the acquisition of the collection of '
          'instruments built up by Paul (who died in 1995) and Edith Egestorff. The '
          'improvement in the collection has also been assisted through a scheme set '
          'up by the Royal Dublin Society, by means of which commercial '
          'organisations sponsored the purchase of important Irish instruments; by '
          'the receipt of instruments from the collection of North Monastery School, '
          'Cork; and by the purchase or donation of instruments discovered during '
          'the Inventory work. The Egestorff acquisitions, the sponsored items, '
          'and the "North Mon" instruments, are identified in the relevant '
          'entries.', 'The collection is to be located in Collins Barracks, '
                      'near the Phoenix Park.', 'Note: brackets around signatures in '
                                                'the above and later entries mean '
                                                'that the signatures have not been '
                                                'seen.', 'INSTRUMENTS RECORDED IN IRISH SALES - SAL',
          'Note: The instruments in this section were seen in shops and in auction', 'sales.',
          'Their whereabouts are now unknown.', 'UNIVERSITY COLLEGE GALWAY ENGINEERING - UGE',
          'Note: This collection is particularly rich in instruments signed by the firm set up by Patrick Adie ('
          '1821-1886), the youngest son of Alexander Adie of Edinburgh.',
          'Patrick worked in London from 1848 until his death, but instruments continued to be signed by the firm '
          'until 1942 - see Clifton 1995,4.',
          'UNIVERSITY COLLEGE GALWAY GEOLOGY - UGG', 'INSTRUMENTS BY THE GRUBB FIRM OF DUBLIN',  'BIRR CASTLE - BIR', 'The Leviathan is due to be restored as the centre-piece of the display at the Historic Science Centre planned for Birr Castle', 'BLACKROCK COLLEGE - BLA', 'CHESTER BEATTY LIBRARY - CBL', 'Note: The Museum is to be re-located in Dublin Castle', 'COMMISSIONERS OF IRISH LIGHTS - CIL', 'DUBLIN CIVIC MUSEUM - DCM', 'DUNSINK OBSERVATORY - DUN', "IRISH DISTILLERS' GROUP LIMITED - IDG", 'METEOROLOGICAL SERVICE - MET', 'INSTRUMENTS IN MISCELLANEOUS INSTITUTIONS - MIS', 'The "MIS" category includes instruments in locations housing five or less items:', 'NATIONAL BOTANIC GARDENS - NBG', 'NEWBRIDGE HOUSE', 'NATIONAL MARITIME MUSEUM - NMM', 'INSTRUMENTS IN PRIVATE COLLECTIONS - PRI', 'Note: The names and addresses of the owners of the instruments in this section are not published, to protect their privacy. However, if anyone wishes to have more information about any of the instruments, the author will contact the owner to see if he/she will allow this and, if so, will pass on the necessary details to the enquirer.', "QUEEN'S UNIVERSITY BELFAST - CHEMISTRY - QBC", 'Note: Several of the items in this collection relate to Thomas Andrews (1813-1885), Professor of Chemistry at what was then Queen\'s College Belfast. He was the first to prove that ozone was a form of oxygen, and he carried out pioneering work on the "critical point" of gases - the temperature above which gases cannot be liquefied.', 'QUEEN\'S UNIVERSITY BELFAST - ENGINEERING - QBE', 'QUEEN\'S UNIVERSITY BELFAST - PHYSICS - QBP', 'ROYAL COLLEGE OF SURGEONS IN IRELAND - RCS', 'Note: The College has a large collection of medical instruments, the listing of which is outside the parameters of this Inventory. However, instruments in this collection, which would normally be included in listings of other scientific instrument collections, have been detailed, these being chiefly microscopes, medical coils, and some thermometers.', 'ROYAL DUBLIN SOCIETY - RDS', 'STATE LABORATORY - STL', 'STROKESTOWN PARK HOUSE - STR', 'TRINITY COLLEGE DUBLIN, CHEMISTRY - TDC', 'TRINITY COLLEGE DUBLIN, ENGINEERING - TDE', 'TRINITY COLLEGE DUBLIN, PHYSICS - TDP', 'UNIVERSITY COLLEGE CORK PHYSICS - UCP', 'UNIVERSITY COLLEGE DUBLIN ENGINEERING - UDE', 'UNIVERSITY COLLEGE DUBLIN PHYSICS - UDP', 'ULSTER FOLK AND TRANSPORT MUSEUM - UFM', 'UNIVERSITY COLLEGE GALWAY PHYSICS - UGP', 'Note: Two catalogue lists of the instruments in the collection of the Galway Physics Department have survived - see Bibliography under Curtis 1861 and Galway 1902.', 'ULSTER MUSEUM - ULS', 'LOCATED OUTSIDE IRELAND', 'Abbreviations used in the listing:', 'INSTRUMENTS LOCATED OUTSIDE IRELAND WITH IRISH SIGNATURES (EXCEPT THOSE SIGNED BY THE GRUBB FIRM)', 'Abbreviations used in the listing:', 'BIBLIOGRAPHY', 'ADAMS, GEORGE (WILLIAM JONES) 1797',
'Note: Many of the instruments in this collection have labels which show the room in which they were originally located, but also letters and numbers (e.g. "F42"), which refer to a note-book catalogue. Although it does not give dates, this catalogue can be most useful in identifying difficult items. It is referred to, when used, as the "Queen\'s catalogue list".']

for r in remove:
    # print(r)
    bolds.remove(r)

pattern = '[A-Z]([a-zA-Z]|[0-9])+[0-9]'
# len(re.findall(pattern, string))
counte = 0
temp_remove = ""
for oth in bolds:
    if len(re.findall(pattern, oth)) == 0:
        print(oth)
        temp_remove = temp_remove + "'{}', ".format(oth)
        counte += 1
print(temp_remove)

bolds_info = {}
found = False
counter = 0
latest = '3950 ARM024 ALIDADE'
for para in document.paragraphs:
    # print('++++++++++++',para.text,'++++++++++++')
    if para.text == "BIBLIOGRAPHY":
        found = False
        break
    # if found and para.text not in remove:
    #     temp_bolds_info = bolds_info.get(latest, [])
    #     temp_bolds_info.append(para.text)
    #     bolds_info[latest] = temp_bolds_info

    for run in para.runs:
        if run.bold and run.text != 'IRISH NATIONAL INVENTORY OF HISTORIC SCIENTIFIC INSTRUMENTS':
            if run.text == bolds[counter]:
                # print("FOUND!")
                latest = run.text
                counter += 1
                found = True
        if found and para.text not in remove:
            temp_bolds_info = bolds_info.get(latest, [])
            temp_bolds_info.append(para.text)
            bolds_info[latest] = temp_bolds_info
            #  and run.text != 'AURORA' and run.text != 'COULOMB'
            # bolds.append(run.text)

# print(counter)
# print(len(bolds))
# all_names, all_maker, all_dims, all_other = [], [], [], []
# for b in bolds:
#     print(b)
#     all_names.append(b)
#     current_info = bolds_info[b]
#     # print(current_info)
#     all_maker.append(current_info[0])
#     all_dims.append(current_info[1])
#     other_info = ''
#     for info in current_info[2:]:
#         other_info = other_info + info
#     all_other.append(other_info)

all_names, all_code, all_areas , all_maker, all_dims, all_other = [], [], [], [], [], []

for ff in bolds[:4486]:
    print(ff,"+++",bolds_info[ff])
    temp_info = bolds_info[ff]
    #all_names.append(temp_info[0])
    temp_name = temp_info[0].split()
    all_code.append(temp_name[0])
    all_areas.append(temp_name[1])

    rest_name = ' '.join(temp_name[2:])
    # for na in temp_name[2:]:
    #     rest_name = rest_name + ' ' +na
    all_names.append(rest_name)

    all_maker.append(temp_info[1])
    if len(temp_info) > 2:
        all_dims.append(temp_info[2])
        rest_info = ''
        for inf in temp_info[3:]:
            rest_info = rest_info + inf
        all_other.append(rest_info)
    else:
        all_dims.append(' ')
        all_other.append(' ')

for ff in bolds[4487:]:
    temp_info = bolds_info[ff]
    # all_names.append(temp_info[0])



    temp_name = temp_info[0].split()
    all_code.append(' ')
    all_areas.append(temp_name[0])

    rest_name = ' '.join(temp_name[1:])
    # for na in temp_name[2:]:
    #     rest_name = rest_name + ' ' +na
    all_names.append(rest_name)

    all_maker.append(temp_info[1])
    if len(temp_info) > 2:
        all_dims.append(temp_info[2])
        rest_info = ''
        for inf in temp_info[3:]:
            rest_info = rest_info + inf
        all_other.append(rest_info)
    else:
        all_dims.append(' ')
        all_other.append(' ')

#print(all_names)
print(len(all_names))
print(len(all_maker))
print(len(all_dims))
print(len(all_other))
for i in range(100):
    print("++",all_names[i], "++",all_maker[i], "++",all_dims[i], "++",all_other[i])

df = pd.DataFrame(list(zip(all_names, all_code, all_areas, all_maker, all_dims, all_other)),
                      columns=['item name', 'item code', 'area code', 'item maker', 'dimensions/date', 'description'])
df.to_csv("1mollan_instruments.csv")
    # if found:
    #     temp_bolds_info = bolds_info.get(latest, [])
    #     temp_bolds_info.append(para.text)
    #     bolds_info[latest] = temp_bolds_info
    # for run in para.runs:
    #     # if run.italic and run.text != 'IRISH NATIONAL INVENTORY OF HISTORIC SCIENTIFIC INSTRUMENTS':
    #     #     italics.append(run.text)
    #     if run.bold and run.text != 'IRISH NATIONAL INVENTORY OF HISTORIC SCIENTIFIC INSTRUMENTS' and run.text != 'AURORA' and run.text != 'COULOMB':
    #         temp_bolds_info = bolds_info.get(latest, [])
    #         temp_bolds_info = temp_bolds_info[:-1]
    #         bolds_info[latest] = temp_bolds_info
    #
    #         bolds.append(run.text)
    #         found = True
    #         counter = 0
    #         latest = run.text

# bolds_info = {}
# found = False
# latest = '3950 ARM024 ALIDADE'
# for para in document.paragraphs:
#     # print('++++++++++++',para.text,'++++++++++++')
#     if found:
#         temp_bolds_info = bolds_info.get(latest, [])
#         temp_bolds_info.append(para.text)
#         bolds_info[latest] = temp_bolds_info
#     for run in para.runs:
#         # if run.italic and run.text != 'IRISH NATIONAL INVENTORY OF HISTORIC SCIENTIFIC INSTRUMENTS':
#         #     italics.append(run.text)
#         if run.bold and run.text != 'IRISH NATIONAL INVENTORY OF HISTORIC SCIENTIFIC INSTRUMENTS' and run.text != 'AURORA' and run.text != 'COULOMB':
#             temp_bolds_info = bolds_info.get(latest, [])
#             temp_bolds_info = temp_bolds_info[:-1]
#             bolds_info[latest] = temp_bolds_info
#
#             bolds.append(run.text)
#             found = True
#             counter = 0
#             latest = run.text


#
#
# all_names, all_maker, all_dims, all_other = [], [], [], []
# for b in bolds:
#     print(b)
#     all_names.append(b)
#     current_info = bolds_info[b]
#     # print(current_info)
#     all_maker.append(current_info[0])
#     all_dims.append(current_info[1])
#     other_info = ''
#     for info in current_info[2:]:
#         other_info = other_info + info
#     all_other.append(other_info)
#
# df = pd.DataFrame(list(zip(all_names, all_maker, all_dims, all_other)),
#                   columns=['item name', 'item maker', 'dimensions/date', 'description'])
# df.to_csv("mollan_instruments.csv")


####### ^^ CODE FOR MAIN SECTION OF MOLLAN ^^ #########
####### VV CODE FOR INDEX SECTION OF MOLLAN VV #########
# document = Document('Mollan_Irish_National_Inventory_of_Historic_Scientific_Instruments_1995.docx')
# paras = []
# runs = []
# for para in document.paragraphs:
#     #print(para.text)
#     paras.append(para.text)
#     for run in para.runs:
#         runs.append(run.text)
#         # print(run.text)
#
# print(paras.index('INDEX BY NAME OF MAKER OR SUPPLIER'))
# paras = paras [47560:]
# # print(paras)
#
# print(runs.index('INDEX BY NAME OF MAKER OR SUPPLIER'))
# runs = runs[35315:]
# #print(runs)
# no_page_numbers_runs = [run for run in paras if len(run)>3 and run != 'IRISH NATIONAL INVENTORY OF HISTORIC SCIENTIFIC INSTRUMENTS']
#
# example = "Abdulla al-A'immah ASTROLABE 4524 CBL001 Abu-t-Tahir Muhammad ASTROLABE QUADRANT 4530 CBL007 AC COIL - INDUCTION, RUHMKORFF 0393 RDS076"
# maker_name = []
# other = []
# for run in no_page_numbers_runs:
#     # print(run)
#     name_index = re.search(' ([A-Z]+/?-?(\(\?\))?)+ ', run).start()
#     maker_name.append(run[:name_index])
#     other.append(run[name_index:])
#
# # for i in range(len(names)):
# #     print(names[i], '++++++', other[i])
#
# inventory_code = []
# location_code = []
# instrument_name = []
# pattern = '[A-Z]([a-zA-Z]|[0-9])+[0-9]'
# # len(re.findall(pattern, string))
# counte = 0
# for oth in other:
#     if len(re.findall(pattern, oth)) > 1:
#         print(oth)
#         counte += 1
#     split_oth = oth.split()
#     inventory_code.append(split_oth[-2])
#     location_code.append(split_oth[-1])
#     temp_instrument_name = ' '.join(split_oth[:-2])
#     instrument_name.append(temp_instrument_name)
#
# print(counte)

####### ^^ CODE FOR INDEX SECTION OF MOLLAN ^^ #########

# for i in range(len(maker_name)):
#      print(maker_name[i], '++++++', instrument_name[i], '++++++', inventory_code[i], '++++++', location_code[i])
#
# df = pd.DataFrame(list(zip(maker_name, instrument_name, inventory_code, location_code)),
#                   columns=['maker name', 'instrument name', 'inventory code', 'location code'])
# df.to_csv("mollan_index_by_name_of_maker.csv")

# regex = re.compile('\b[A-Z]+\b')
# testt = re.search(' [A-Z]+ ', run).start()
# # testt = regex.search("Abdulla al-A'immah ASTROLABE 4524 CBL001 Abu-t-Tahir Muhammad ASTROLABE QUADRANT 4530 CBL007 AC COIL - INDUCTION, RUHMKORFF 0393 RDS076").group(1)
# print(example[testt:])


# '3985 ARM059 DIVIDED CIRCLE WITH REVOLVING DISC'
# for para in document.paragraphs:
#     print('++++++++++++',para.text,'++++++++++++')
#     for run in para.runs:
#         # if run.italic and run.text != 'IRISH NATIONAL INVENTORY OF HISTORIC SCIENTIFIC INSTRUMENTS':
#         #     italics.append(run.text)
#         if run.bold and run.text != 'IRISH NATIONAL INVENTORY OF HISTORIC SCIENTIFIC INSTRUMENTS':
#             bolds.append(run.text)

# from PyPDF2 import PdfFileWriter, PdfFileReader
# input_pdf = PdfFileReader("Instruments of science.pdf")
# output = PdfFileWriter()
# for i in range(136):
#     output.addPage(input_pdf.getPage(i+400))
# with open("Instruments-of-science-400-end.pdf", "wb") as output_stream:
#     output.write(output_stream)

# def get_runs_paras_from_docx(document_name):
#     document = Document(document_name)
#     paras = []
#     runs = []
#     for para in document.paragraphs:
#         #print(para.text)
#         paras.append(para.text)
#         for run in para.runs:
#             runs.append(run.text)
#     return runs, paras
#
# runss, parass = get_runs_paras_from_docx("CMC masters-apprentices copy.docx")
# for run in parass:
#     print(run)


####### EXTRACT FROM CMC MASTERS APPRENTICES ########
# from tika import parser # pip install tika
#
# raw = parser.from_file('CMC masters-apprentices copy.pdf')
# print(raw['content'])
#
# with open("new-CMC-masters-apprentices.txt", "w") as text_file:
#     text_file.write(raw['content'])